from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Mapping

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from sqlalchemy.orm import Session

from app.core.config import Settings, get_settings
from app.models import NationalEconomyClassificationCase


SCENARIO = "national_economy_classification"
PENDING_STATUS = "pending_classification"

FIELD_LABELS = {
    "enterprise_name": "企业名称",
    "unified_social_credit_code": "统一社会信用代码",
    "business_scope": "营业执照经营范围（全文）",
    "main_business": "主营业务",
    "main_business_revenue_share": "主营业务及营收占比",
    "core_products_services": "核心产品 / 服务名称",
    "loan_purpose": "贷款用途详细描述",
    "counterparty_name": "贸易合同本次交易对手名称",
    "counterparty_business_industry": "交易对手主营业务 / 所属行业",
    "trade_goods_services": "贸易合同核心交易品类 / 服务内容",
    "industry_chain_position": "企业产业链定位",
    "industry_position_competitiveness": "企业行业定位与核心竞争力",
    "credit_approval_opinion": "授信审批意见",
}


@dataclass(frozen=True)
class TemplateValidationIssues:
    missing: tuple[str, ...] = ()
    duplicate: tuple[str, ...] = ()
    unrecognized: tuple[str, ...] = ()


class NationalEconomyTemplateError(ValueError):
    def __init__(self, issues: TemplateValidationIssues) -> None:
        self.issues = issues
        details = []
        if issues.missing:
            details.append(f"缺失标签: {', '.join(issues.missing)}")
        if issues.duplicate:
            details.append(f"重复标签: {', '.join(issues.duplicate)}")
        if issues.unrecognized:
            details.append(f"无法识别标签: {', '.join(issues.unrecognized)}")
        super().__init__("; ".join(details))


def read_template_bytes(settings: Settings | None = None) -> bytes:
    template_path = (settings or get_settings()).national_economy_template_path
    return template_path.read_bytes()


def parse_template(document_bytes: bytes) -> dict[str, str]:
    return parse_template_fields(document_bytes, FIELD_LABELS)


def parse_template_fields(
    document_bytes: bytes,
    field_labels: Mapping[str, str],
    field_aliases: Mapping[str, tuple[str, ...]] | None = None,
) -> dict[str, str]:
    try:
        document = Document(BytesIO(document_bytes))
    except (PackageNotFoundError, ValueError, KeyError) as exc:
        raise NationalEconomyTemplateError(
            TemplateValidationIssues(unrecognized=("文件不是可解析的 .docx 模板",))
        ) from exc

    aliases = field_aliases or {}
    label_to_field = {
        accepted_label: field
        for field, label in field_labels.items()
        for accepted_label in (label, *aliases.get(field, ()))
    }
    values: dict[str, str] = {}
    duplicate_labels: list[str] = []
    unrecognized_labels: list[str] = []

    def add_value(label: str, value: str) -> None:
        normalized_label = label.strip()
        field = label_to_field.get(normalized_label)
        if field is None:
            unrecognized_labels.append(normalized_label)
            return
        if field in values:
            duplicate_labels.append(normalized_label)
            return
        values[field] = value.strip()

    table_rows_found = False
    for table in document.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            label = row.cells[0].text.strip()
            if not label or label == "字段名称":
                continue
            table_rows_found = True
            add_value(label, row.cells[1].text)

    # Keep accepting the original paragraph-based template. New table-based
    # documents may contain titles and instructions, so their paragraphs are
    # intentionally ignored once labeled table rows are present.
    if not table_rows_found:
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            label, separator, value = text.partition("：")
            if not separator:
                unrecognized_labels.append(text)
                continue
            add_value(label, value)

    missing_labels = [
        label for field, label in field_labels.items() if field not in values
    ]
    issues = TemplateValidationIssues(
        missing=tuple(missing_labels),
        duplicate=tuple(duplicate_labels),
        unrecognized=tuple(unrecognized_labels),
    )
    if issues.missing or issues.duplicate or issues.unrecognized:
        raise NationalEconomyTemplateError(issues)

    return {field: values[field] for field in field_labels}


def create_case_from_template(
    session: Session,
    document_bytes: bytes,
    original_filename: str,
) -> NationalEconomyClassificationCase:
    input_payload = parse_template(document_bytes)
    case = NationalEconomyClassificationCase(
        scenario=SCENARIO,
        input_payload=input_payload,
        original_filename=Path(original_filename).name,
        status=PENDING_STATUS,
    )
    session.add(case)
    session.commit()
    session.refresh(case)
    return case

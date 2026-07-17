from datetime import datetime, timezone
from io import BytesIO
from unittest.mock import MagicMock

import pytest
from docx import Document
from app.models import NationalEconomyClassificationCase

from app.services.five_articles_case_ingestion import (
    create_five_articles_case_from_template,
    parse_five_articles_template,
)
from app.services.scenario_case_handlers import FIVE_ARTICLES_CASE_HANDLER
from app.services.national_economy_case_ingestion import (
    NationalEconomyTemplateError,
)
from app.services.scenario_registry import (
    AGRICULTURE_RELATED_REGISTRATION,
    DIGITAL_FINANCE_REGISTRATION,
    GREEN_FINANCE_REGISTRATION,
    MULTI_SCENARIO_FINANCE_REGISTRATIONS,
    PENSION_FINANCE_REGISTRATION,
    ScenarioRegistration,
)


def _filled_formal_agriculture_template() -> bytes:
    document = Document(
        BytesIO(AGRICULTURE_RELATED_REGISTRATION.template_path().read_bytes())
    )
    for row, field in zip(
        document.tables[0].rows[1:],
        AGRICULTURE_RELATED_REGISTRATION.field_schema,
        strict=True,
    ):
        row.cells[1].text = f"正式模板值-{field.key}"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _three_column_template(
    profile: ScenarioRegistration,
    *,
    issue: str | None = None,
) -> bytes:
    rows = [
        [field.label, f"测试值-{field.key}"]
        for field in profile.field_schema
    ]
    if issue == "missing":
        rows.pop()
    elif issue == "duplicate":
        rows.append(rows[0].copy())
    elif issue == "unrecognized":
        rows[0][0] = "无法识别的字段"

    document = Document()
    document.add_heading(f"{profile.name}企业信息采集表", level=1)
    document.add_paragraph("请在第二列填写，第三列为填写提示。")
    table = document.add_table(rows=1, cols=3)
    for cell, value in zip(
        table.rows[0].cells,
        ("字段名称", "填写内容", "填写提示"),
        strict=True,
    ):
        cell.text = value
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        cells[2].text = "不得进入案例字段"

    output = BytesIO()
    document.save(output)
    return output.getvalue()


@pytest.mark.parametrize("profile", MULTI_SCENARIO_FINANCE_REGISTRATIONS)
def test_formal_three_column_template_creates_profile_owned_case(
    profile: ScenarioRegistration,
) -> None:
    session = MagicMock()

    case = create_five_articles_case_from_template(
        session,
        profile.template_path().read_bytes(),
        f"../{profile.name}案例.docx",
        profile,
    )

    assert case.scenario == profile.id
    assert case.original_filename == f"{profile.name}案例.docx"
    assert case.status == "pending_classification"
    assert list(case.input_payload) == [field.key for field in profile.field_schema]
    assert len(case.input_payload) == len(profile.field_schema)
    assert "不得进入案例字段" not in case.input_payload.values()
    session.add.assert_called_once_with(case)
    session.commit.assert_called_once_with()
    session.refresh.assert_called_once_with(case)


def test_previous_pension_template_without_new_optional_share_field_is_accepted() -> None:
    document = Document(
        BytesIO(PENSION_FINANCE_REGISTRATION.template_path().read_bytes())
    )
    table = document.tables[0]
    optional_row = next(
        row
        for row in table.rows
        if row.cells[0].text.strip()
        == "该笔贷款实际投向养老产业占总贷款额度比"
    )
    table._tbl.remove(optional_row._tr)
    output = BytesIO()
    document.save(output)

    payload = parse_five_articles_template(
        output.getvalue(),
        PENSION_FINANCE_REGISTRATION,
    )

    assert payload["pension_loan_direction_share"] == ""
    assert list(payload) == [
        field.key for field in PENSION_FINANCE_REGISTRATION.field_schema
    ]


@pytest.mark.parametrize(
    ("profile", "issue", "issue_name"),
    [
        (GREEN_FINANCE_REGISTRATION, "missing", "missing"),
        (DIGITAL_FINANCE_REGISTRATION, "duplicate", "duplicate"),
        (PENSION_FINANCE_REGISTRATION, "unrecognized", "unrecognized"),
    ],
)
def test_invalid_profile_template_reports_issue_without_creating_case(
    profile: ScenarioRegistration,
    issue: str,
    issue_name: str,
) -> None:
    session = MagicMock()

    with pytest.raises(NationalEconomyTemplateError) as error:
        create_five_articles_case_from_template(
            session,
            _three_column_template(profile, issue=issue),
            f"{profile.id}-{issue}.docx",
            profile,
        )

    assert getattr(error.value.issues, issue_name)
    session.add.assert_not_called()
    session.commit.assert_not_called()
    session.refresh.assert_not_called()


@pytest.mark.parametrize("profile", MULTI_SCENARIO_FINANCE_REGISTRATIONS)
def test_three_column_parser_ignores_hint_column(
    profile: ScenarioRegistration,
) -> None:
    payload = parse_five_articles_template(
        _three_column_template(profile),
        profile,
    )

    assert payload == {
        field.key: f"测试值-{field.key}"
        for field in profile.field_schema
    }


def test_formal_agriculture_template_creates_case_with_all_twenty_fields() -> None:
    session = MagicMock()

    case = create_five_articles_case_from_template(
        session,
        _filled_formal_agriculture_template(),
        "../涉农企业.docx",
        AGRICULTURE_RELATED_REGISTRATION,
    )

    assert case.scenario == AGRICULTURE_RELATED_REGISTRATION.id
    assert case.status == "pending_classification"
    assert list(case.input_payload) == [
        field.key for field in AGRICULTURE_RELATED_REGISTRATION.field_schema
    ]
    assert len(case.input_payload) == 20
    assert case.input_payload["project_content"] == "正式模板值-project_content"
    session.add.assert_called_once_with(case)
    session.commit.assert_called_once_with()
    session.refresh.assert_called_once_with(case)


@pytest.mark.parametrize("issue", ["missing", "duplicate", "unrecognized"])
def test_agriculture_template_issues_do_not_create_case(issue: str) -> None:
    session = MagicMock()

    with pytest.raises(NationalEconomyTemplateError) as error:
        create_five_articles_case_from_template(
            session,
            _three_column_template(AGRICULTURE_RELATED_REGISTRATION, issue=issue),
            f"涉农-{issue}.docx",
            AGRICULTURE_RELATED_REGISTRATION,
        )

    assert getattr(error.value.issues, issue)
    session.add.assert_not_called()
    session.commit.assert_not_called()
    session.refresh.assert_not_called()


def test_agriculture_case_detail_uses_all_profile_fields() -> None:
    case = NationalEconomyClassificationCase(
        id=42,
        scenario=AGRICULTURE_RELATED_REGISTRATION.id,
        input_payload={
            field.key: f"详情值-{field.key}"
            for field in AGRICULTURE_RELATED_REGISTRATION.field_schema
        },
        original_filename="涉农企业.docx",
        status="pending_classification",
        created_at=datetime.now(timezone.utc),
        updated_at=datetime.now(timezone.utc),
    )

    response = FIVE_ARTICLES_CASE_HANDLER.case_response(
        case, AGRICULTURE_RELATED_REGISTRATION
    )

    assert [(field.field, field.label) for field in response.input_fields] == [
        (field.key, field.label)
        for field in AGRICULTURE_RELATED_REGISTRATION.field_schema
    ]
    assert len(response.input_fields) == 20
    assert response.input_fields[-1].value == "详情值-credit_approval_opinion"

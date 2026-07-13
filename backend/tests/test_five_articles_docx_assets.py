from copy import deepcopy
from pathlib import Path

import pytest
from docx import Document

from scripts.preflight_five_articles_docx_assets import (
    ASSET_SPECS,
    DocxAssetValidationError,
    STAGE_A_FIELD_ALIASES,
    validate_docx_asset,
)


def _spec(scenario_id: str):
    return next(spec for spec in ASSET_SPECS if spec.scenario_id == scenario_id)


@pytest.mark.parametrize("spec", ASSET_SPECS, ids=lambda spec: spec.scenario_id)
def test_formal_docx_asset_matches_locked_field_contract(spec) -> None:
    labels = validate_docx_asset(spec.path, spec)

    assert len(labels) == len(spec.fields)
    assert len(labels) == len(set(labels))
    assert all(set(labels).intersection(aliases) for aliases in STAGE_A_FIELD_ALIASES.values())


def test_previous_pension_layout_fails_when_fields_are_embedded_in_hints(
    tmp_path: Path,
) -> None:
    spec = _spec("pension_finance")
    document = Document(spec.path)
    table = document.tables[0]

    project_row = next(
        row for row in table.rows if row.cells[0].text.strip() == "项目建设 / 运营内容"
    )
    project_previous = project_row._element.getprevious()
    project_previous.tc_lst[2].text = (
        "文本，固定资产贷款、项目贷款必须填写） "
        "项目建设 / 运营内容：（长文本，项目的具体建设内容、产出效益，项目类贷款必填"
    )
    table._tbl.remove(project_row._element)

    trade_row = next(
        row for row in table.rows if row.cells[0].text.strip() == "核心交易品类 / 服务内容"
    )
    trade_previous = trade_row._element.getprevious()
    trade_previous.tc_lst[2].text = (
        "文本，用于验证交易真实性与行业产业链定位） "
        "核心交易品类 / 服务内容：（文本，本次交易的具体货物、服务品类"
    )
    table._tbl.remove(trade_row._element)

    broken_path = tmp_path / "养老金融模版-修正前.docx"
    document.save(broken_path)

    with pytest.raises(DocxAssetValidationError) as error:
        validate_docx_asset(broken_path, spec)

    message = str(error.value)
    assert "字段行应为 18，实际 16" in message
    assert "项目建设 / 运营内容" in message
    assert "核心交易品类 / 服务内容" in message


@pytest.mark.parametrize("defect", ("missing", "duplicate", "embedded"))
def test_docx_asset_preflight_rejects_each_field_row_defect(
    defect: str,
    tmp_path: Path,
) -> None:
    spec = _spec("digital_finance")
    document = Document(spec.path)
    table = document.tables[0]
    target = next(
        row for row in table.rows if row.cells[0].text.strip() == "项目建设 / 运营内容"
    )

    if defect == "missing":
        table._tbl.remove(target._element)
    elif defect == "duplicate":
        table._tbl.append(deepcopy(target._element))
    else:
        target.cells[2].text += " 核心交易品类 / 服务内容：不得嵌入提示"

    broken_path = tmp_path / f"digital-{defect}.docx"
    document.save(broken_path)

    with pytest.raises(DocxAssetValidationError):
        validate_docx_asset(broken_path, spec)

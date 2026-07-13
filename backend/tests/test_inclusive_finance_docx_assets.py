from copy import deepcopy
from pathlib import Path
from unittest.mock import MagicMock

import pytest
from docx import Document

from app.services.five_articles_case_ingestion import (
    create_five_articles_case_from_template,
    parse_five_articles_template,
)
from app.services.national_economy_case_ingestion import (
    FIELD_LABELS,
    NationalEconomyTemplateError,
)
from app.services.scenario_registry import INCLUSIVE_FINANCE_REGISTRATION


TEMPLATE_PATH = Path("模板文件/五篇大文章/普惠金融模版.docx")
FARMER_FIELD_KEYS = {
    "farmer_long_term_town_resident",
    "farmer_town_village_resident",
    "farmer_nonlocal_resident_over_one_year",
    "farmer_state_farm_employee_or_rural_individual_business",
}
REQUIRED_DETERMINATION_FIELD_KEYS = {
    "entity_type",
    "enterprise_scale_type",
    "total_assets",
    "annual_revenue",
    "employee_count",
    "credit_amount",
    "credit_variety",
    "loan_purpose",
    "project_name",
    "project_content",
    *FARMER_FIELD_KEYS,
}


def _field_rows(document: Document) -> list[tuple[str, str, str]]:
    tables = [table for table in document.tables if len(table.columns) == 3]
    assert len(tables) == 1
    table = tables[0]
    assert tuple(cell.text.strip() for cell in table.rows[0].cells) == (
        "字段名称",
        "填写内容",
        "填写提示",
    )
    return [
        tuple(cell.text.strip() for cell in row.cells)
        for row in table.rows[1:]
        if any(cell.text.strip() for cell in row.cells)
    ]


def _expected_template_labels() -> set[str]:
    return {
        field.aliases[0]
        if field.key in {"enterprise_name", "counterparty_name", "trade_goods_services"}
        else field.label
        for field in INCLUSIVE_FINANCE_REGISTRATION.field_schema
    }


def _validate_inclusive_docx_asset(path: Path) -> None:
    rows = _field_rows(Document(path))
    labels = tuple(row[0] for row in rows)

    assert len(rows) == 31
    assert len(labels) == len(set(labels))
    assert set(labels) == _expected_template_labels()
    assert not {
        label
        for _, _, hint in rows
        for label in labels
        if f"{label}：" in hint or f"{label}:" in hint
    }


def test_inclusive_finance_docx_asset_matches_exact_31_field_schema() -> None:
    rows = _field_rows(Document(TEMPLATE_PATH))
    labels = tuple(row[0] for row in rows)
    schema = INCLUSIVE_FINANCE_REGISTRATION.field_schema
    schema_keys = {field.key for field in schema}

    _validate_inclusive_docx_asset(TEMPLATE_PATH)
    assert len(schema) == 31
    assert len(schema_keys) == len(schema)
    assert INCLUSIVE_FINANCE_REGISTRATION.stage_a_field_keys == tuple(
        key for key in FIELD_LABELS if key != "main_business"
    )
    assert "main_business" not in schema_keys
    assert REQUIRED_DETERMINATION_FIELD_KEYS <= schema_keys


def test_real_inclusive_template_creates_case_without_template_validation_issues() -> None:
    session = MagicMock()

    case = create_five_articles_case_from_template(
        session,
        TEMPLATE_PATH.read_bytes(),
        "../普惠金融案例.docx",
        INCLUSIVE_FINANCE_REGISTRATION,
    )

    assert case.scenario == "inclusive_finance"
    assert case.original_filename == "普惠金融案例.docx"
    assert list(case.input_payload) == [
        field.key for field in INCLUSIVE_FINANCE_REGISTRATION.field_schema
    ]
    assert len(case.input_payload) == 31
    assert "main_business" not in case.input_payload
    session.add.assert_called_once_with(case)
    session.commit.assert_called_once_with()
    session.refresh.assert_called_once_with(case)


@pytest.mark.parametrize("defect", ("missing", "duplicate", "embedded"))
def test_inclusive_finance_docx_asset_rejects_each_field_row_defect(
    defect: str,
    tmp_path: Path,
) -> None:
    document = Document(TEMPLATE_PATH)
    table = document.tables[0]
    target = next(
        row for row in table.rows if row.cells[0].text.strip() == "主体类型"
    )

    if defect == "missing":
        table._tbl.remove(target._element)
    elif defect == "duplicate":
        table._tbl.append(deepcopy(target._element))
    else:
        target.cells[2].text += " 企业规模类型：不得嵌入填写提示"

    broken_path = tmp_path / f"inclusive-{defect}.docx"
    document.save(broken_path)

    with pytest.raises(AssertionError):
        _validate_inclusive_docx_asset(broken_path)

    if defect != "embedded":
        with pytest.raises(NationalEconomyTemplateError) as error:
            parse_five_articles_template(
                broken_path.read_bytes(),
                INCLUSIVE_FINANCE_REGISTRATION,
            )

        assert getattr(error.value.issues, defect)

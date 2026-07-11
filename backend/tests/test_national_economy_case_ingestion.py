from io import BytesIO
from pathlib import Path
from unittest.mock import MagicMock

import pytest
from docx import Document

from app.core.config import Settings
from app.services.national_economy_case_ingestion import (
    FIELD_LABELS,
    NationalEconomyTemplateError,
    create_case_from_template,
    parse_template,
    read_template_bytes,
)


FIXTURES = Path(__file__).parent / "fixtures" / "national_economy"


def _fixture_bytes(name: str) -> bytes:
    return (FIXTURES / name).read_bytes()


def test_valid_template_creates_pending_case_with_thirteen_fields() -> None:
    session = MagicMock()

    case = create_case_from_template(
        session,
        _fixture_bytes("valid.docx"),
        "../企业A.docx",
    )

    assert case.scenario == "national_economy_classification"
    assert case.status == "pending_classification"
    assert case.original_filename == "企业A.docx"
    assert list(case.input_payload) == list(FIELD_LABELS)
    assert case.input_payload["enterprise_name"] == "南京示例科技有限公司"
    assert case.input_payload["credit_approval_opinion"] == "同意授信"
    session.add.assert_called_once_with(case)
    session.commit.assert_called_once_with()
    session.refresh.assert_called_once_with(case)


def test_unfilled_fields_are_saved_as_empty_strings() -> None:
    payload = parse_template(_fixture_bytes("empty_fields.docx"))

    assert len(payload) == 13
    assert payload["enterprise_name"] == "南京示例科技有限公司"
    assert payload["loan_purpose"] == ""
    assert payload["credit_approval_opinion"] == ""


def test_table_template_is_parsed_and_ignores_title_and_instructions() -> None:
    document = Document()
    document.add_heading("国民经济行业分类企业信息采集表", level=1)
    document.add_paragraph("请在第二列填写企业真实经营信息，第三列为填写提示。")
    table = document.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "字段名称"
    table.rows[0].cells[1].text = "填写内容"
    table.rows[0].cells[2].text = "填写提示"
    for field, label in FIELD_LABELS.items():
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = "测试企业" if field == "enterprise_name" else ""
        cells[2].text = "请按实际情况填写"
    output = BytesIO()
    document.save(output)

    payload = parse_template(output.getvalue())

    assert list(payload) == list(FIELD_LABELS)
    assert payload["enterprise_name"] == "测试企业"
    assert payload["credit_approval_opinion"] == ""


@pytest.mark.parametrize(
    ("fixture_name", "issue_name", "problem_label"),
    [
        ("missing_label.docx", "missing", "授信审批意见"),
        ("duplicate_label.docx", "duplicate", "企业名称"),
        ("unrecognized_label.docx", "unrecognized", "企业全称"),
    ],
)
def test_invalid_labels_report_issues_without_creating_case(
    fixture_name: str,
    issue_name: str,
    problem_label: str,
) -> None:
    session = MagicMock()

    with pytest.raises(NationalEconomyTemplateError) as error:
        create_case_from_template(session, _fixture_bytes(fixture_name), fixture_name)

    assert problem_label in getattr(error.value.issues, issue_name)
    session.add.assert_not_called()
    session.commit.assert_not_called()
    session.refresh.assert_not_called()


def test_template_download_returns_original_docx_bytes(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    original_bytes = _fixture_bytes("valid.docx")
    template_path.write_bytes(original_bytes)
    settings = Settings(
        _env_file=None,
        national_economy_template_path=template_path,
    )

    assert read_template_bytes(settings) == original_bytes
